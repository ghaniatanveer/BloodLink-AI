from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 11)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title_run = title.add_run(
        "BloodLink-AI: A Role-Based Blood Donation Management and Intelligent Assistance Platform"
    )
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    author = doc.add_paragraph(
        "Project Report (IEEE Style)\n"
        "Prepared for: BloodLink-AI Development Submission\n"
        f"Date: {date.today().strftime('%B %d, %Y')}"
    )
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, "Abstract", 1)
    add_body(
        doc,
        "BloodLink-AI is a full-stack web-based blood donation management platform that connects donors and hospitals "
        "through a role-based workflow, city-wise discovery, blood-group compatibility logic, and an AI-supported "
        "assistant for donation guidance. The system was implemented using Node.js, Express.js, SQLite (better-sqlite3), "
        "JWT authentication, and a responsive web frontend. The architecture supports account registration/login, "
        "hospital-side blood request publishing, donor-side request discovery and acceptance, hospital approval and completion "
        "tracking, donor/hospital histories, and a resilient chatbot that prioritizes command-based responses and provides "
        "graceful fallback responses under external LLM quota failures. The project emphasizes practical local deployment, "
        "zero-cost database operation, robust request handling, and clear user experience for emergency blood coordination."
    )

    add_heading(doc, "Keywords", 1)
    add_body(
        doc,
        "Blood Donation; Role-Based Access Control; Express.js; SQLite; JWT; Healthcare Workflow; Chatbot; "
        "Gemini API; Groq API; Fallback Architecture."
    )

    add_heading(doc, "I. Introduction", 1)
    add_body(
        doc,
        "Blood availability is a time-sensitive public health requirement. Existing donation processes are often fragmented "
        "across social networks and phone calls, causing delays in emergency matching. BloodLink-AI addresses this by offering "
        "a structured digital workflow for donors and hospitals in a single platform. The project was designed to be practical "
        "for local deployment, easy to run on a low-resource machine, and extensible for future production hardening."
    )

    add_heading(doc, "II. Problem Statement and Objectives", 1)
    add_body(
        doc,
        "The core problem is the lack of a streamlined, role-aware blood request system that can quickly match compatible donors "
        "within the same city while preserving a clear approval workflow for hospitals. The objectives were:"
    )
    add_bullets(
        doc,
        [
            "Implement secure donor/hospital authentication and authorization.",
            "Build city-aware and blood-group-aware donor-request matching.",
            "Create a complete lifecycle: request creation -> donor acceptance -> hospital decision -> completion.",
            "Maintain donation/request history for both user roles.",
            "Provide AI assistance for workflow guidance and donation education.",
            "Ensure graceful behavior even when external AI APIs fail or are rate-limited.",
        ],
    )

    add_heading(doc, "III. System Overview", 1)
    add_body(
        doc,
        "BloodLink-AI follows a client-server architecture. The frontend (HTML/CSS/JS) consumes REST APIs from an Express backend. "
        "Data persistence is implemented using SQLite through better-sqlite3. JWT is used for stateless authentication. "
        "Role checks (donor/hospital) and in-memory rate limiting protect endpoint usage."
    )
    add_body(doc, "Main modules include:")
    add_bullets(
        doc,
        [
            "Authentication module (`/api/auth`)",
            "Request workflow module (`/api/requests`)",
            "Directory search module (`/api/directory`)",
            "Chat assistant module (`/api/chat`)",
            "Static web interface (`public/` assets)",
        ],
    )

    add_heading(doc, "IV. Technology Stack", 1)
    add_bullets(
        doc,
        [
            "Runtime: Node.js (ES Modules)",
            "Backend Framework: Express.js",
            "Database: SQLite (better-sqlite3)",
            "Authentication: JWT (`jsonwebtoken`) + Password hashing (`bcryptjs`)",
            "AI Integrations: Google Gemini (`@google/generative-ai`) and Groq (`groq-sdk`)",
            "Environment Management: dotenv",
            "Development Tooling: nodemon",
            "Frontend: Vanilla HTML, CSS, JavaScript",
        ],
    )

    add_heading(doc, "V. Database Design", 1)
    add_body(
        doc,
        "The database schema includes `users`, `requests`, `request_candidates`, and `meta` tables. "
        "Key constraints and indexes support role consistency, request status progression, and query performance."
    )
    add_bullets(
        doc,
        [
            "`users`: stores donor/hospital identity, contact, city, role, blood group, location, and credentials.",
            "`requests`: stores hospital-created blood requests with urgency, units, and status fields.",
            "`request_candidates`: junction table for donor acceptance and hospital action states.",
            "`meta`: stores seeding flags/version metadata.",
            "Indexes on request city/status, hospital request ownership, and users role/city.",
        ],
    )
    add_body(
        doc,
        "A deterministic seeding strategy populates 30 donor records per city across a broad city list, with generated unique "
        "names, emails, and local-format phone numbers."
    )

    add_heading(doc, "VI. Functional Workflow", 1)
    add_heading(doc, "A. Registration and Login", 2)
    add_body(
        doc,
        "Users register as donor or hospital via `/api/auth/register`. Input validation enforces required fields, role, password "
        "length, and blood-group validity for donors. Passwords are hashed before storage. `/api/auth/login` verifies credentials "
        "and issues JWT tokens."
    )

    add_heading(doc, "B. Hospital Request Lifecycle", 2)
    add_bullets(
        doc,
        [
            "Hospital creates request: `/api/requests/hospital/create`",
            "Hospital lists own requests: `/api/requests/hospital/list`",
            "Hospital approves/rejects donor candidate: `/api/requests/hospital/approve/:id`",
            "Hospital marks completion: `/api/requests/hospital/complete/:id`",
            "Hospital accesses completed donor history: `/api/requests/hospital/history`",
        ],
    )

    add_heading(doc, "C. Donor Interaction", 2)
    add_bullets(
        doc,
        [
            "Donor views nearby compatible open/in-progress requests by city: `/api/requests/donor/nearby`",
            "Donor accepts a request: `/api/requests/donor/accept/:id`",
            "Donor views participation history: `/api/requests/donor/history`",
        ],
    )
    add_body(
        doc,
        "Compatibility checks are enforced using blood-group logic (`canDonateTo`) and city matching. "
        "Optional distance in kilometers is computed when both donor and hospital coordinates are available."
    )

    add_heading(doc, "D. Public Directory Search", 2)
    add_body(
        doc,
        "The directory module provides public searchable lists for donors and hospitals via `/api/directory/donors` and "
        "`/api/directory/hospitals`, enabling discovery-oriented UX."
    )

    add_heading(doc, "VII. AI Assistant Design", 1)
    add_body(
        doc,
        "The chat subsystem supports command shortcuts and general donation Q&A. Command triggers such as "
        "\"find donors\" and \"open requests\" return deterministic database-backed responses. For broader conversational "
        "queries, the system tries Groq and/or Gemini models depending on available API keys."
    )
    add_body(doc, "Reliability mechanisms include:")
    add_bullets(
        doc,
        [
            "Model fallback list for Gemini (`GEMINI_MODEL_FALLBACKS`).",
            "429 quota retry-delay parsing and controlled wait.",
            "Conversation trimming to limit prompt size.",
            "Offline canned educational responses for common donation topics.",
            "Guaranteed friendly fallback response instead of raw provider error dumps.",
        ],
    )

    add_heading(doc, "VIII. Security and Validation", 1)
    add_bullets(
        doc,
        [
            "JWT-based route protection (`requireAuth`).",
            "Role-based authorization (`requireRole`).",
            "Password hashing with bcrypt.",
            "Basic in-memory rate limiting (60 requests/minute per IP).",
            "Input validation for registration and core entities.",
            "Global error handling and route-level 404 handling.",
        ],
    )

    add_heading(doc, "IX. API Summary", 1)
    add_body(doc, "Representative endpoint groups:")
    add_bullets(
        doc,
        [
            "Auth: `/api/auth/register`, `/api/auth/login`, `/api/auth/me`",
            "Requests: donor and hospital workflow endpoints under `/api/requests/*`",
            "Directory: `/api/directory/donors`, `/api/directory/hospitals`",
            "Chat: `/api/chat/send`, `/api/chat/conversations`, `/api/chat/conversation/:id`",
        ],
    )

    add_heading(doc, "X. User Interface and Experience", 1)
    add_body(
        doc,
        "The frontend is optimized for search-first interaction similar to practical blood lookup portals. "
        "It supports role-oriented actions, city filtering, quick-chat commands, and English-only user-facing content."
    )

    add_heading(doc, "XI. Testing and Verification", 1)
    add_body(doc, "The implemented verification approach includes:")
    add_bullets(
        doc,
        [
            "Manual endpoint-level testing for auth, request lifecycle, and histories.",
            "Role mismatch and invalid input checks (401/403/400 pathways).",
            "Chat command tests for deterministic outputs.",
            "AI failure-path testing under quota-limited scenarios.",
            "Server start/port verification and database readiness checks.",
        ],
    )

    add_heading(doc, "XII. Results", 1)
    add_body(
        doc,
        "The final system meets the stated functional goals: role-based user onboarding, city and compatibility constrained "
        "matching, full donation workflow tracking, history visibility, and resilient assistant behavior. The migration from "
        "MongoDB to SQLite reduced setup complexity and supports zero-cost local deployment without external database dependencies."
    )

    add_heading(doc, "XIII. Limitations and Future Work", 1)
    add_bullets(
        doc,
        [
            "Current rate limiter is in-memory; distributed deployment would require centralized storage.",
            "No formal automated test suite is included yet.",
            "Conversation memory is process-local and not persisted across restarts.",
            "Production hardening may add HTTPS termination, audit logs, and secret rotation workflows.",
            "Future scope includes notifications (SMS/WhatsApp), geospatial indexing, and analytics dashboards.",
        ],
    )

    add_heading(doc, "XIV. Conclusion", 1)
    add_body(
        doc,
        "BloodLink-AI demonstrates a practical, deployable, and extensible blood donation platform that combines operational "
        "workflow logic with user assistance capabilities. The architecture is suitable for academic evaluation and local pilots, "
        "while remaining ready for incremental production-grade enhancements."
    )

    add_heading(doc, "References", 1)
    refs = [
        "[1] Express.js Documentation. [Online]. Available: https://expressjs.com/",
        "[2] SQLite Documentation. [Online]. Available: https://www.sqlite.org/docs.html",
        "[3] JSON Web Token Introduction. [Online]. Available: https://jwt.io/introduction",
        "[4] bcryptjs Package Documentation. [Online]. Available: https://www.npmjs.com/package/bcryptjs",
        "[5] Google Generative AI for Node.js. [Online]. Available: https://www.npmjs.com/package/@google/generative-ai",
        "[6] Groq SDK Documentation. [Online]. Available: https://www.npmjs.com/package/groq-sdk",
        "[7] IEEE Author Center. [Online]. Available: https://journals.ieeeauthorcenter.ieee.org/",
    ]
    for r in refs:
        add_body(doc, r)

    output_path = "d:\\Projects\\BloodLink-AI\\BloodLink_AI_IEEE_Report.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
